from insert import * 
from parse_sn1per import *
from parse_utils import *
from docx.shared import Cm
from docx.text.run import Run

def add_subdomains(doc, workspace):
    subddomains_list =  parse_subdomains(workspace)

    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).text = "П/п"
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).text  = "Поддомены"
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].text = str(i)
        row_cells[0].width = Cm(3)

        row_cells[1].text = domain
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "Поддомены")

def add_harvester_asn(doc, workspace):
    subddomains_list =  parse_the_harvester_asn(workspace)

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).text = "П/п"
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).text  = "ASN"
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].text = str(i)
        row_cells[0].width = Cm(3)

        row_cells[1].text = domain
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "ASN")

def add_harvester_interesting_url(doc, workspace):
    subddomains_list =  parse_the_harvester_interesting_url(workspace)

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).text = "П/п"
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).text  = "Популярные URL"
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].text = str(i)
        row_cells[0].width = Cm(3)

        row_cells[1].text = domain
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "Популярные URL")

def add_forms(doc, workspace):
    subddomains_list =  parse_forms(workspace)

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).text = "П/п"
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).text  = "URL форм"
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].text = str(i)
        row_cells[0].width = Cm(3)

        row_cells[1].text = domain
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "Формы")

def add_emails(doc, workspace):
    subddomains_list =  parse_emails(workspace)

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).text = "П/п"
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).text  = "Адрес почты"
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].text = str(i)
        row_cells[0].width = Cm(3)

        row_cells[1].text = domain
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "Адреса электронной почты")

def add_public_files(doc, workspace):
    subddomains_list =  parse_public_files(workspace)

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).text = "П/п"
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).text  = "URL файлов"
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].text = str(i)
        row_cells[0].width = Cm(3)

        row_cells[1].text = domain
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "Файлы в публичном доступе")

def add_tech(doc, workspace):
    subddomains_list =  parse_tech(workspace)

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).text = "П/п"
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).text  = "Название фреймворка, технологии"
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].text = str(i)
        row_cells[0].width = Cm(3)

        row_cells[1].text = domain
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "Применяемые фреймворки, технологии")

def add_robots(doc, workspace):
    
    subddomains_list =  parse_robots(workspace)

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).text = "П/п"
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).text  = "Строка файла"
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].text = str(i)
        row_cells[0].width = Cm(3)

        row_cells[1].text = domain
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "Содержимое robots.txt")

