
from .insert import *
from .parse_utils import *
from docx.shared import Cm
import itertools
from docx.text.run import Run
import copy

def add_masscan(doc, workspace):
    
    data =  parse_masscan(workspace)
    
    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).paragraphs[0].text='П/п'
    table.cell(0,0).paragraphs[0].style='Таблица'
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).paragraphs[0].text  = "TCP-порт"
    table.cell(0,1).paragraphs[0].style='Таблица'
    table.cell(0,1).width = Cm(13)
    # table.cell(0,1).text  = "IP-адресс"
    # table.cell(0,1).width = Cm(6.5)
    # table.cell(0,2).text  = "Порт-TCP"
    # table.cell(0,2).width = Cm(6.5)
    i = 1
    for ip in data ['ip_list_sorted']:
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[1])

        
        row_cells[0].paragraphs[0].text = ip
        row_cells[0].paragraphs[0].style = "Таблица"
        run = isolate_run(row_cells[0].paragraphs[0], 0, len(ip))
        # run=row_cells[0].paragraphs[0].add_run()
        run.font.bold=True

        row_cells[0].width = Cm(16)
        i = 1
        for port in data ['data'] [ip]:
            row_cells = table.add_row().cells

            row_cells[0].paragraphs[0].text = str(i)
            row_cells[0].paragraphs[0].style = "Таблица"
            row_cells[0].width = Cm(3)

            # row_cells[1].text = ip
            # row_cells[1].width = Cm(6.5)

            row_cells[1].paragraphs[0].text = str(port)
            row_cells[1].paragraphs[0].style = "Таблица"
            row_cells[1].width = Cm(13)
            i += 1
    
    move_table_after(doc, table, "Открытые TCP-порты")

def add_nmap(doc, workspace):
    
    data =  parse_nmap(workspace)


    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).paragraphs[0].text='П/п'
    table.cell(0,0).paragraphs[0].style='Таблица'
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).paragraphs[0].text  = "TCP-порт"
    table.cell(0,1).paragraphs[0].style='Таблица'
    table.cell(0,1).width = Cm(13)
    # table.cell(0,1).text  = "IP-адресс"
    # table.cell(0,1).width = Cm(6.5)
    # table.cell(0,2).text  = "Порт-TCP"
    # table.cell(0,2).width = Cm(6.5)

    i = 1   
    for host in data ["nmaprun"]["host"]:
        if  host.get("ports") != None:
            row_cells = table.add_row().cells
            row_cells[0].merge(row_cells[1])
            ip = ""
            #print(type(host["address"]))
            if type(host["address"]) == list:
                ip = host ["address"][0]["@addr"]
            else: ip = host ["address"]["@addr"]

            row_cells[0].paragraphs[0].text = ip
            row_cells[0].paragraphs[0].style = "Таблица"
            run = isolate_run(row_cells[0].paragraphs[0], 0, len(ip))
            # run=row_cells[0].paragraphs[0].add_run()
            run.font.bold=True

            row_cells[0].width = Cm(16)
            i = 1
            # print (host["ports"])
            for port in host["ports"] ["port"]:
                row_cells = table.add_row().cells

                row_cells[0].paragraphs[0].text = str(i)
                row_cells[0].paragraphs[0].style = "Таблица"
                row_cells[0].width = Cm(3)

                # row_cells[1].text = ip
                # row_cells[1].width = Cm(6.5)

                row_cells[1].paragraphs[0].text = str(port["@portid"])
                row_cells[1].paragraphs[0].style = "Таблица"
                row_cells[1].width = Cm(13)
                i += 1
        else: continue 
    
    move_table_after(doc, table, "Открытые TCP-порты")


def add_harvester_subdomains(doc, workspace):

    subddomains_list =  parse_harvester(workspace) ['hosts']

    table = doc.add_table(rows=1, cols=3)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).paragraphs[0].text = "П/п"
    table.cell(0,0).paragraphs[0].style='Таблица'
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).paragraphs[0].text  = "Поддомены"
    table.cell(0,1).paragraphs[0].style='Таблица'
    table.cell(0,1).width = Cm(6.5)
    table.cell(0,2).paragraphs[0].text  = "IP-адреса"
    table.cell(0,2).paragraphs[0].style='Таблица'
    table.cell(0,2).width = Cm(6.5)
    i = 1
    for domain_ip in subddomains_list:
        splited_str = domain_ip.split(":",1)
        domain = splited_str [0]
        if len(splited_str) == 1:
            ip = ''
            continue
        else:
            ip = splited_str [1]

        row_cells = table.add_row().cells

        row_cells[0].paragraphs[0].text = str(i)
        row_cells[0].paragraphs[0].style='Таблица'
        row_cells[0].width = Cm(3)

        row_cells[1].paragraphs[0].text = domain
        row_cells[1].paragraphs[0].style='Таблица'
        row_cells[1].width = Cm(6.5)

        
        row_cells[2].paragraphs[0].text = ip
        row_cells[2].paragraphs[0].style='Таблица'
        row_cells[2].width = Cm(6.5)
        i += 1
    
    move_table_after(doc, table, "Поддомены")

def add_harvester_asn(doc, workspace):
    subddomains_list =  parse_harvester(workspace) ['asns']

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).paragraphs[0].text = "П/п"
    table.cell(0,0).paragraphs[0].style='Таблица'
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).paragraphs[0].text  = "ASN"
    table.cell(0,1).paragraphs[0].style='Таблица'
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].paragraphs[0].text = str(i)
        row_cells[0].paragraphs[0].style='Таблица'
        row_cells[0].width = Cm(3)

        row_cells[1].paragraphs[0].text = domain
        row_cells[1].paragraphs[0].style='Таблица'
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "ASN")

def add_harvester_interesting_url(doc, workspace):
    subddomains_list =  parse_harvester(workspace) ['interesting_urls']

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).paragraphs[0].text = "П/п"
    table.cell(0,0).paragraphs[0].style='Таблица'
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).paragraphs[0].text  = "Популярные URL"
    table.cell(0,1).paragraphs[0].style='Таблица'
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].paragraphs[0].text = str(i)
        row_cells[0].paragraphs[0].style='Таблица'
        row_cells[0].width = Cm(3)

        row_cells[1].paragraphs[0].text = domain
        row_cells[1].paragraphs[0].style='Таблица'
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "Популярные URL")

def add_emails(doc, workspace):
    subddomains_list =  parse_harvester(workspace) ['emails']

    # print(subddomains_list)
    table = doc.add_table(rows=1, cols=2)
    table.style = "table"
    table.autofit = False
    table.cell(0,0).paragraphs[0].text = "П/п"
    table.cell(0,0).paragraphs[0].style='Таблица'
    table.cell(0,0).width = Cm(3)
    table.cell(0,1).paragraphs[0].text  = "Email"
    table.cell(0,1).paragraphs[0].style='Таблица'
    table.cell(0,1).width = Cm(13)
    i = 1
    for domain in subddomains_list:
        row_cells = table.add_row().cells

        row_cells[0].paragraphs[0].text = str(i)
        row_cells[0].paragraphs[0].style='Таблица'
        row_cells[0].width = Cm(3)

        row_cells[1].paragraphs[0].text = domain
        row_cells[1].paragraphs[0].style='Таблица'
        row_cells[1].width = Cm(13)
        i += 1
    
    move_table_after(doc, table, "Адреса электронной почты")

def isolate_run(paragraph, start, end):
    """Return docx.text.Run object containing only `paragraph.text[start:end]`.

    Runs are split as required to produce a new run at the `start` that ends at `end`.
    Runs are unchanged if the indicated range of text already occupies its own run. The
    resulting run object is returned.

    `start` and `end` are as in Python slice notation. For example, the first three
    characters of the paragraph have (start, end) of (0, 3). `end` is not the index of
    the last character. These correspond to `match.start()` and `match.end()` of a regex
    match object and `s[start:end]` of Python slice notation.
    """
    rs = tuple(paragraph._p.r_lst)

    def advance_to_run_containing_start(start, end):
        """Return (r_idx, start, end) triple indicating start run and adjusted offsets.

        The start run is the run the `start` offset occurs in. The returned `start` and
        `end` values are adjusted to be relative to the start of `r_idx`.
        """
        # --- add 0 at end so `r_ends[-1] == 0` ---
        r_ends = tuple(itertools.accumulate(len(r.text) for r in rs)) + (0,)
        r_idx = 0
        while start >= r_ends[r_idx]:
            r_idx += 1
        skipped_rs_offset = r_ends[r_idx - 1]
        return rs[r_idx], r_idx, start - skipped_rs_offset, end - skipped_rs_offset

    def split_off_prefix(r, start, end):
        """Return adjusted `end` after splitting prefix off into separate run.

        Does nothing if `r` is already the start of the isolated run.
        """
        if start > 0:
            prefix_r = copy.deepcopy(r)
            r.addprevious(prefix_r)
            r.text = r.text[start:]
            prefix_r.text = prefix_r.text[:start]
        return end - start

    def split_off_suffix(r, end):
        """Split `r` at `end` such that suffix is in separate following run."""
        suffix_r = copy.deepcopy(r)
        r.addnext(suffix_r)
        r.text = r.text[:end]
        suffix_r.text = suffix_r.text[end:]

    def lengthen_run(r, r_idx, end):
        """Add prefixes of following runs to `r` until `end` is reached."""
        while len(r.text) < end:
            suffix_len_reqd = end - len(r.text)
            r_idx += 1
            next_r = rs[r_idx]
            if len(next_r.text) <= suffix_len_reqd:
                # --- subsume next run ---
                r.text = r.text + next_r.text
                next_r.getparent().remove(next_r)
                continue
            if len(next_r.text) > suffix_len_reqd:
                # --- take prefix from next run ---
                r.text = r.text + next_r.text[:suffix_len_reqd]
                next_r.text = next_r.text[suffix_len_reqd:]

    r, r_idx, start, end = advance_to_run_containing_start(start, end)
    end = split_off_prefix(r, start, end)

    # --- if run is longer than isolation-range we need to split-off a suffix run ---
    if len(r.text) > end:
        split_off_suffix(r, end)
    # --- if run is shorter than isolation-range we need to lengthen it by taking text
    # --- from subsequent runs
    elif len(r.text) < end:
        lengthen_run(r, r_idx, end)

    return Run(r, paragraph)